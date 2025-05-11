from docx import Document

# Membuat dokumen Word
doc = Document()
doc.add_heading('PROPOSAL PROYEK', 0)
doc.add_heading('Monitoring Detak Jantung dan Kadar Oksigen Menggunakan Arduino Uno', level=1)

# 1. Latar Belakang
doc.add_heading('1. Latar Belakang', level=2)
doc.add_paragraph(
    "Kesehatan jantung dan kadar oksigen dalam darah merupakan parameter penting dalam pemantauan kondisi tubuh, "
    "terutama pada pasien dengan gangguan pernapasan, jantung, atau yang sedang dalam pemulihan medis. Di era digital "
    "saat ini, teknologi mikrokontroler seperti Arduino Uno memungkinkan kita untuk merancang alat monitoring kesehatan "
    "yang murah, portabel, dan dapat diakses oleh masyarakat luas. Dengan menggabungkan sensor MAX30100 atau MAX30102, "
    "alat ini mampu mengukur detak jantung (BPM) dan kadar oksigen dalam darah (SpO₂) secara real time."
)

# 2. Tujuan
doc.add_heading('2. Tujuan', level=2)
doc.add_paragraph('- Membangun alat monitoring detak jantung dan kadar oksigen berbasis Arduino Uno.')
doc.add_paragraph('- Menampilkan data BPM dan SpO₂ melalui layar OLED dan Serial Monitor.')
doc.add_paragraph('- Memberikan solusi monitoring kesehatan yang murah, efektif, dan portabel.')

# 3. Alat dan Bahan
doc.add_heading('3. Alat dan Bahan', level=2)
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'No'
hdr_cells[1].text = 'Nama Komponen'
hdr_cells[2].text = 'Jumlah'

components = [
    ("1", "Arduino Uno", "1 buah"),
    ("2", "Sensor MAX30100 / MAX30102", "1 buah"),
    ("3", "OLED Display SSD1306 (I2C)", "1 buah"),
    ("4", "Breadboard", "1 buah"),
    ("5", "Kabel Jumper", "Beberapa"),
    ("6", "Laptop + Arduino IDE", "1 set"),
    ("7", "Powerbank / Sumber daya USB", "1 buah"),
]

for item in components:
    row_cells = table.add_row().cells
    row_cells[0].text = item[0]
    row_cells[1].text = item[1]
    row_cells[2].text = item[2]

# 4. Metode Kerja
doc.add_heading('4. Metode Kerja', level=2)
doc.add_paragraph('a. Rangkaian Sistem', style='List Bullet')
doc.add_paragraph('- Menghubungkan sensor MAX30100 ke Arduino Uno melalui pin I2C: SDA ke A4, SCL ke A5.')
doc.add_paragraph('- Menghubungkan OLED SSD1306 juga ke pin I2C yang sama.')
doc.add_paragraph('- Menggunakan breadboard dan jumper untuk koneksi listrik + logika.')

doc.add_paragraph('b. Pemrograman', style='List Bullet')
doc.add_paragraph('- Menggunakan Arduino IDE untuk menulis dan mengunggah kode ke Arduino Uno.')
doc.add_paragraph('- Kode memanfaatkan library: MAX30100_PulseOximeter, Adafruit GFX, dan Adafruit SSD1306.')
doc.add_paragraph('- Data BPM dan SpO₂ akan ditampilkan di OLED dan Serial Monitor secara real time.')

doc.add_paragraph('c. Pengujian', style='List Bullet')
doc.add_paragraph('- Meletakkan jari di atas sensor MAX30100.')
doc.add_paragraph('- Membandingkan hasil pembacaan alat dengan alat oximeter medis jika tersedia (validasi).')
doc.add_paragraph('- Menganalisis kestabilan dan akurasi pembacaan.')

# 5. Hasil yang Diharapkan
doc.add_heading('5. Hasil yang Diharapkan', level=2)
doc.add_paragraph('- Alat dapat menampilkan detak jantung dan kadar oksigen secara akurat.')
doc.add_paragraph('- Dapat bekerja secara real time dan portabel.')
doc.add_paragraph('- Menjadi prototipe awal sistem pemantauan kesehatan berbasis mikrokontroler.')

# 6. Penutup
doc.add_heading('6. Penutup', level=2)
doc.add_paragraph(
    "Dengan proyek ini, diharapkan kita dapat menciptakan perangkat pemantau kesehatan sederhana yang berguna dalam "
    "kehidupan sehari-hari, terutama untuk pemantauan mandiri di rumah. Proyek ini juga menjadi langkah awal dalam "
    "pengembangan sistem kesehatan berbasis IoT yang dapat terintegrasi ke aplikasi seluler atau cloud di masa depan."
)

# Simpan file
file_path = "Proposal_Monitoring_Detak_Jantung_Oksigen_Arduino.docx"
doc.save(file_path)
print(f"File proposal berhasil disimpan: {file_path}")
